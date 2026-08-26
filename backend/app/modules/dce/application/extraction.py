"""Deterministic, private DCE document extraction without HTTP exposure or AI."""

from __future__ import annotations

import logging
from collections.abc import Iterable
from dataclasses import dataclass
from datetime import UTC, datetime
from hashlib import sha256
from io import BytesIO
from typing import Protocol
from uuid import UUID, uuid5
from zipfile import BadZipFile, ZipFile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.orm import Session, sessionmaker

from app.modules.dce.application.commands import (
    DceExtractionFragmentInput,
    RecordDceDocumentExtractionCommand,
)
from app.modules.dce.infrastructure.models.dce_staging import DceStagedObjectRecord
from app.modules.dce.infrastructure.models.dce_version import DceDocumentRecord, DceVersionRecord
from app.platform.events.dispatcher import CommandContext, CommandDispatcher, DispatchResult

logger = logging.getLogger(__name__)

EXTRACTOR_ID = "smart-ao-deterministic"
EXTRACTOR_VERSION = "1"
MAX_SOURCE_BYTES = 128 * 1024 * 1024
MAX_PDF_PAGES = 2_000
MAX_DOCX_PARAGRAPHS = 100_000
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MAX_XLSX_SHEETS = 200
MAX_XLSX_TEXT_CELLS = 500_000
MAX_TEXT_LINES = 500_000
MAX_FRAGMENT_CHARS = 8_000
MAX_TOTAL_CHARS = 10_000_000
MAX_TOTAL_DOCUMENT_CHARS = 20_000_000
SYSTEM_EXTRACTION_ACTOR_ID = UUID("00000000-0000-0000-0000-000000000013")


class PrivateDocumentStoragePort(Protocol):
    """Private server-side read port; no route can obtain a storage key through it."""

    async def read_bytes(self, *, storage_key: str, max_bytes: int) -> bytes: ...


class AdvancedDocumentExtractionPort(Protocol):
    """Optional local parser port; it must return only bounded source fragments."""

    def extract(self, *, media_type: str, source_bytes: bytes) -> ExtractionProjection: ...


@dataclass(frozen=True, slots=True)
class ExtractedFragment:
    ordinal: int
    locator_json: dict[str, object]
    text: str


@dataclass(frozen=True, slots=True)
class ExtractionProjection:
    status: str
    failure_code: str | None
    fragments: tuple[ExtractedFragment, ...]
    extractor_id: str = EXTRACTOR_ID
    extractor_version: str = EXTRACTOR_VERSION


class ExtractionLimitError(ValueError):
    """Expected terminal limit breach; it must not expose parser details."""


class DceDocumentExtractionService:
    """Reads one consumed original and records only its bounded deterministic projection."""

    def __init__(
        self,
        *,
        session_factory: sessionmaker[Session],
        dispatcher: CommandDispatcher,
        storage: PrivateDocumentStoragePort,
        advanced_extractor: AdvancedDocumentExtractionPort | None = None,
    ) -> None:
        self._session_factory = session_factory
        self._dispatcher = dispatcher
        self._storage = storage
        self._advanced_extractor = advanced_extractor

    async def extract(
        self,
        *,
        tenant_id: UUID,
        dce_document_id: UUID,
        now: datetime | None = None,
    ) -> DispatchResult:
        effective_now = now or datetime.now(tz=UTC)
        document = self._load_private_document(
            tenant_id=tenant_id,
            dce_document_id=dce_document_id,
        )
        if document.byte_size > MAX_SOURCE_BYTES:
            projection = ExtractionProjection(
                status="REJECTED_LIMIT",
                failure_code="EXTRACTION_LIMIT",
                fragments=(),
            )
        else:
            try:
                source_bytes = await self._storage.read_bytes(
                    storage_key=document.storage_key,
                    max_bytes=MAX_SOURCE_BYTES,
                )
            except (FileNotFoundError, OSError, ValueError):
                projection = ExtractionProjection(
                    status="FAILED_SAFE",
                    failure_code="PRIVATE_DOCUMENT_UNAVAILABLE",
                    fragments=(),
                )
            else:
                if (
                    len(source_bytes) != document.byte_size
                    or sha256(source_bytes).hexdigest() != document.sha256
                ):
                    projection = ExtractionProjection(
                        status="FAILED_SAFE",
                        failure_code="PRIVATE_DOCUMENT_INTEGRITY_MISMATCH",
                        fragments=(),
                    )
                else:
                    projection = _project_document(
                        media_type=document.media_type,
                        source_bytes=source_bytes,
                        advanced_extractor=self._advanced_extractor,
                    )
        command = _recording_command(
            document=document,
            projection=projection,
        )
        return self._dispatcher.dispatch(
            command=command,
            context=CommandContext(
                tenant_id=tenant_id,
                actor_id=SYSTEM_EXTRACTION_ACTOR_ID,
                actor_kind="SYSTEM",
                received_at=effective_now,
            ),
        )

    def _load_private_document(
        self,
        *,
        tenant_id: UUID,
        dce_document_id: UUID,
    ) -> DceDocumentRecord:
        with self._session_factory() as session:
            document = session.scalar(
                select(DceDocumentRecord).where(
                    DceDocumentRecord.tenant_id == tenant_id,
                    DceDocumentRecord.id == dce_document_id,
                )
            )
            if document is None:
                raise ValueError("DCE_DOCUMENT_REQUIRED")
            dce_version = session.scalar(
                select(DceVersionRecord).where(
                    DceVersionRecord.tenant_id == tenant_id,
                    DceVersionRecord.id == document.dce_version_id,
                )
            )
            staged_object = session.scalar(
                select(DceStagedObjectRecord).where(
                    DceStagedObjectRecord.tenant_id == tenant_id,
                    DceStagedObjectRecord.id == document.storage_object_id,
                )
            )
            if dce_version is None or dce_version.lifecycle not in {"ADMITTED", "SUPERSEDED"}:
                raise ValueError("DCE_VERSION_NOT_ADMITTED")
            if dce_version.integrity != "VERIFIED":
                raise ValueError("DCE_VERSION_NOT_VERIFIED")
            if (
                staged_object is None
                or staged_object.state != "CONSUMED"
                or staged_object.consumed_by_dce_version_id != document.dce_version_id
            ):
                raise ValueError("DOCUMENT_STORAGE_NOT_CONSUMED")
            session.expunge(document)
            return document


def _recording_command(
    *,
    document: DceDocumentRecord,
    projection: ExtractionProjection,
) -> RecordDceDocumentExtractionCommand:
    extraction_identity = uuid5(
        document.id,
        f"{document.sha256.lower()}:{projection.extractor_id}:{projection.extractor_version}",
    )
    return RecordDceDocumentExtractionCommand(
        command_id=extraction_identity,
        idempotency_key=extraction_identity,
        correlation_id=document.dce_version_id,
        extraction_id=extraction_identity,
        dce_document_id=document.id,
        input_sha256=document.sha256.lower(),
        extractor_id=projection.extractor_id,
        extractor_version=projection.extractor_version,
        status=projection.status,
        extracted_char_count=sum(len(fragment.text) for fragment in projection.fragments),
        failure_code=projection.failure_code,
        fragments=[
            DceExtractionFragmentInput(
                ordinal=fragment.ordinal,
                locator_json=fragment.locator_json,
                text=fragment.text,
                text_sha256=sha256(fragment.text.encode("utf-8")).hexdigest(),
            )
            for fragment in projection.fragments
        ],
    )


def _project_document(
    *,
    media_type: str,
    source_bytes: bytes,
    advanced_extractor: AdvancedDocumentExtractionPort | None = None,
) -> ExtractionProjection:
    try:
        try:
            fragments = tuple(_extract_fragments(media_type=media_type, source_bytes=source_bytes))
        except ValueError as error:
            if str(error) != "MEDIA_TYPE_UNSUPPORTED":
                raise
            fragments = ()
        if fragments:
            return ExtractionProjection(
                status="COMPLETED",
                failure_code=None,
                fragments=fragments,
            )
        if advanced_extractor is not None:
            advanced_projection = advanced_extractor.extract(
                media_type=media_type,
                source_bytes=source_bytes,
            )
            if advanced_projection.status != "UNSUPPORTED":
                return advanced_projection
        return ExtractionProjection(
            status="FAILED_SAFE",
            failure_code="EMPTY_EXTRACTED_TEXT",
            fragments=(),
        )
    except ExtractionLimitError:
        return ExtractionProjection(
            status="REJECTED_LIMIT",
            failure_code="EXTRACTION_LIMIT",
            fragments=(),
        )
    except ValueError as error:
        if str(error) == "MEDIA_TYPE_UNSUPPORTED":
            return ExtractionProjection(
                status="UNSUPPORTED",
                failure_code="MEDIA_TYPE_UNSUPPORTED",
                fragments=(),
            )
        logger.warning(
            "dce_extraction_value_error",
            extra={"error_type": type(error).__name__, "media_type": media_type},
        )
        return ExtractionProjection(
            status="FAILED_SAFE",
            failure_code="EXTRACTION_PARSE_FAILED",
            fragments=(),
        )
    except Exception as error:
        logger.warning(
            "dce_extraction_parse_failed",
            extra={"error_type": type(error).__name__, "media_type": media_type},
        )
        return ExtractionProjection(
            status="FAILED_SAFE",
            failure_code="EXTRACTION_PARSE_FAILED",
            fragments=(),
        )


def _extract_fragments(*, media_type: str, source_bytes: bytes) -> Iterable[ExtractedFragment]:
    if media_type == "application/pdf":
        return _extract_pdf(source_bytes=source_bytes)
    if media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(source_bytes=source_bytes)
    if media_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        return _extract_xlsx(source_bytes=source_bytes)
    if media_type == "text/plain":
        return _extract_text(source_bytes=source_bytes)
    raise ValueError("MEDIA_TYPE_UNSUPPORTED")


def _extract_pdf(*, source_bytes: bytes) -> tuple[ExtractedFragment, ...]:
    reader = PdfReader(BytesIO(source_bytes))
    if reader.is_encrypted:
        raise ValueError("encrypted PDF")
    if len(reader.pages) > MAX_PDF_PAGES:
        raise ExtractionLimitError
    return _fragmentize(
        (
            ({"kind": "pdf_page", "page": page_number}, page.extract_text() or "")
            for page_number, page in enumerate(reader.pages, start=1)
        )
    )


def _extract_docx(*, source_bytes: bytes) -> tuple[ExtractedFragment, ...]:
    try:
        with ZipFile(BytesIO(source_bytes)) as archive:
            if sum(info.file_size for info in archive.infolist()) > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ExtractionLimitError
            total_bytes = 0
            for info in archive.infolist():
                if info.is_dir():
                    continue
                with archive.open(info) as member:
                    while chunk := member.read(1024 * 1024):
                        total_bytes += len(chunk)
                        if total_bytes > MAX_DOCX_UNCOMPRESSED_BYTES:
                            raise ExtractionLimitError
    except BadZipFile as error:
        raise ValueError("invalid DOCX archive") from error

    document = Document(BytesIO(source_bytes))
    if len(document.paragraphs) > MAX_DOCX_PARAGRAPHS:
        raise ExtractionLimitError
    return _fragmentize(
        (
            ({"kind": "docx_paragraph", "paragraph": index}, paragraph.text)
            for index, paragraph in enumerate(document.paragraphs, start=1)
        )
    )


def _extract_xlsx(*, source_bytes: bytes) -> tuple[ExtractedFragment, ...]:
    workbook = load_workbook(
        BytesIO(source_bytes),
        read_only=True,
        data_only=True,
        keep_links=False,
    )
    try:
        if len(workbook.worksheets) > MAX_XLSX_SHEETS:
            raise ExtractionLimitError
        return _fragmentize(_iter_xlsx_entries(workbook))
    finally:
        workbook.close()


def _iter_xlsx_entries(workbook) -> Iterable[tuple[dict[str, object], str]]:
    cell_count = 0
    for worksheet in workbook.worksheets:
        for row in worksheet.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str) or not cell.value.strip():
                    continue
                cell_count += 1
                if cell_count > MAX_XLSX_TEXT_CELLS:
                    raise ExtractionLimitError
                yield (
                    {
                        "kind": "xlsx_cell",
                        "sheet": worksheet.title,
                        "cell": cell.coordinate,
                    },
                    cell.value,
                )


def _extract_text(*, source_bytes: bytes) -> tuple[ExtractedFragment, ...]:
    text = source_bytes.decode("utf-8", errors="strict")
    lines = text.splitlines()
    if len(lines) > MAX_TEXT_LINES:
        raise ExtractionLimitError
    return _fragmentize(
        (({"kind": "text_line", "line": index}, line) for index, line in enumerate(lines, start=1))
    )


def _fragmentize(entries: Iterable[tuple[dict[str, object], str]]) -> tuple[ExtractedFragment, ...]:
    fragments: list[ExtractedFragment] = []
    total_chars = 0
    for locator, raw_text in entries:
        if len(raw_text) > MAX_TOTAL_CHARS:
            raise ExtractionLimitError
        total_chars += len(raw_text)
        if total_chars > MAX_TOTAL_DOCUMENT_CHARS:
            raise ExtractionLimitError
        normalized = raw_text.strip()
        if not normalized:
            continue
        for part, text in enumerate(_split_text(normalized), start=1):
            fragment_locator = dict(locator)
            fragment_locator["part"] = part
            fragments.append(
                ExtractedFragment(
                    ordinal=len(fragments) + 1,
                    locator_json=fragment_locator,
                    text=text,
                )
            )
    return tuple(fragments)


def _split_text(text: str) -> Iterable[str]:
    for start in range(0, len(text), MAX_FRAGMENT_CHARS):
        yield text[start : start + MAX_FRAGMENT_CHARS]
